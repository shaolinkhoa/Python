# -*- coding: utf-8 -*-
"""
Created on Mon Aug 20 15:18:05 2018

@author: ukn1hc
"""

from docx import Document

# Input / read exist document
document = Document('test_tab_indence.docx')

paragraphs = document.paragraphs

# Add shift + tab with \t
p = document.add_paragraph('one\ttwo\tthree')


#bold
#p.add_run('Name:\t\t').bold = True



#Read a paragraph
a = paragraphs[0].text

# Add ctrl+enter with \n
p = document.add_paragraph('NEW_LINE_HERE_2\nTest_with_key_code\nShould_in_para')
# or Add with add_break
p = document.add_paragraph('NEW_LINE_HERE_3')
run = p.add_run()
run.add_break()
p.add_run('and some')
p.add_run().add_break()
p.add_run('down')

'''
paragraphs[9].text 
>>Sub_signal1
paragraphs[10].text
>>Khoa

We want to add a paragraph between [9]&[10]
we have to use paragraphs[10].insert_paragraph_before()
'''
#cursor_para = paragraphs[10]
#cursor_para.insert_paragraph_before('insert before para[10]: Khoa')


# Search word
'''
for paragraph in document.paragraphs:
    if 'sea' in paragraph.text:
        print paragraph.text
        paragraph.text = 'new text containing ocean'
'''

#Search in table
'''
for table in document.tables:
    for cell in table.cells:
        for paragraph in cell.paragraphs:
            if 'sea' in paragraph.text:
               ...
'''

# save to new file
# Docx can't export rtf
document.save('demo.docx')